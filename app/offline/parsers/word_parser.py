from __future__ import annotations

from pathlib import Path

from app.utils.text_utils import clean_text


def parse_docx(path: Path) -> dict:
    try:
        from docx import Document  # type: ignore

        doc = Document(path)
        headings: list[str] = []
        paragraphs: list[str] = []
        tables: list[list[str]] = []
        for paragraph in doc.paragraphs:
            text = clean_text(paragraph.text)
            if not text:
                continue
            paragraphs.append(text)
            if paragraph.style and "Heading" in paragraph.style.name:
                headings.append(text)
        for table in doc.tables:
            headers = [clean_text(cell.text) for cell in table.rows[0].cells]
            tables.append(headers)
        raw_text = "\n".join(paragraphs)
        return {
            "pages": [{"page_no": 1, "raw_text": raw_text, "text_blocks": paragraphs}],
            "headings": headings,
            "table_headers": [header for row in tables for header in row if header],
            "raw_text": raw_text,
        }
    except Exception:
        text = clean_text(path.read_text(encoding="utf-8", errors="ignore"))
        return {
            "pages": [{"page_no": 1, "raw_text": text, "text_blocks": [text]}],
            "headings": text.splitlines()[:5],
            "table_headers": [],
            "raw_text": text,
        }
